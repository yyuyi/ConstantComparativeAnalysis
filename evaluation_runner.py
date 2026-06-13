from __future__ import annotations

import argparse
from difflib import SequenceMatcher
import html
import json
import os
import re
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, Iterable, List, Optional, Tuple

BASE_DIR = Path(__file__).resolve().parent

try:
    from . import config
    from .agents.sdk import AgentSDK
except Exception:
    import config  # type: ignore
    from agents.sdk import AgentSDK  # type: ignore


MACHINE_EVALUATION_SCHEMA = (
    "{\"overall_score\": float, \"machine_summary\": str, \"criteria\": "
    "[{\"name\": str, \"score\": int, \"rationale\": str, \"evidence\": [str], "
    "\"risks\": [str], \"recommendations\": [str]}], \"strengths\": [str], "
    "\"limitations\": [str], \"human_review_priorities\": [str]}"
)

SCORING_SCALE = {
    1: "Absent or not auditable: the output provides little/no evidence for this CCA criterion, or the evidence directly contradicts the criterion.",
    2: "Weak: the output gestures toward the criterion but is mostly superficial, inconsistent, poorly evidenced, or hard to audit.",
    3: "Adequate/mixed: the output shows the criterion in some places, but with notable gaps, uneven coverage, or unresolved ambiguity.",
    4: "Strong: the output substantially satisfies the criterion with concrete evidence, with only minor gaps or localized weaknesses.",
    5: "Excellent: the output consistently and clearly satisfies the criterion across the relevant files, with strong evidence and no important unresolved gap.",
}

CRITERION_SCORE_ANCHORS = {
    "Constant comparison process": {
        1: "No auditable incident-to-incident comparison; output mainly lists themes/codes without comparing cases.",
        2: "Some comparison language appears, but it is generic, local only, or not tied to specific incidents.",
        3: "Some incident comparisons are substantive, but cross-case/non-local comparison or reconciliation is uneven.",
        4: "Clear incident-to-incident and cross-case comparison with useful reconciliation; minor gaps remain.",
        5: "Consistent local and non-local comparison across the run; reconciliation clearly consolidates labels while preserving variation and boundary cases.",
    },
    "Traceability and quote evidence": {
        1: "Claims cannot be traced to segment IDs or source text; quotes are absent, fabricated, or not auditable.",
        2: "Some segment IDs or quotes are present, but many are missing, dirty, mismatched, or disconnected from source context.",
        3: "Most major claims have traceable segment IDs/quotes, but quote cleanliness, context fit, or evidence coverage is uneven.",
        4: "Claims are usually supported by clean quotes and segment IDs, with minor traceability or context gaps.",
        5: "Category claims are consistently traceable to clean quotes, segment IDs, and source context; quoted evidence substantively supports the interpretation.",
    },
    "Boundary and negative-case handling": {
        1: "No boundary/negative cases are identified, and category limits are not auditable.",
        2: "Boundary/negative cases are mentioned rarely or superficially, without changing category interpretation.",
        3: "Most major categories include boundary/negative material, but its analytic effect is inconsistent or unclear.",
        4: "Categories usually include grounded boundary/negative cases and explain how they qualify the category.",
        5: "Every final category has a grounded boundary/negative case or explicit no-case rationale, and these cases actively shape category boundaries.",
    },
    "Category differentiation and non-overmerge": {
        1: "Categories are absent, incoherent, or collapsed into very broad themes with little differentiation.",
        2: "Categories exist but are redundant, too thin, or overly broad macro buckets.",
        3: "Categories are partly differentiated, but some overlap, overmerge, or granularity problems remain.",
        4: "Categories are mostly distinct, mid-grained, and grounded in different mechanisms or conditions.",
        5: "Category structure is consistently distinct, appropriately granular, and avoids overmerging while preserving meaningful variation.",
    },
    "Memo-to-synthesis coherence": {
        1: "Memos/synthesis are absent or disconnected from categories and evidence.",
        2: "Memos or synthesis mostly summarize topics without developing comparative insight or relationships.",
        3: "Some memo insights carry into synthesis, but links among incidents, categories, boundaries, and claims are uneven.",
        4: "Synthesis follows from memos and categories, explains relationships, and is mostly evidence-bounded.",
        5: "Memos, category relationships, tensions, boundary cases, and final synthesis form a coherent evidence-bounded analytic chain.",
    },
    "Cross-coder divergence integration": {
        1: "No meaningful cross-coder integration is present for a two-coder run, or one coder dominates without explanation.",
        2: "Coder differences are listed but not interpreted or used to revise category boundaries.",
        3: "Some convergence/divergence is integrated, but boundary impact or minority interpretations are uneven.",
        4: "Integration explains major convergences/divergences and usually shows how disagreement affects categories.",
        5: "Integration consistently preserves both coders' contributions, explains divergence, and uses disagreement to refine category boundaries and synthesis.",
    },
}

CRITERIA = [
    {
        "name": "Constant comparison process",
        "human_prompt": "Does the output show incident-to-incident comparison across local and non-local cases, with useful reconciliation rather than isolated thematic coding?",
    },
    {
        "name": "Boundary and negative-case handling",
        "human_prompt": "Does every final category identify boundary/negative cases or explain why none were found, and do those cases affect interpretation?",
    },
    {
        "name": "Category differentiation and non-overmerge",
        "human_prompt": "Are categories mid-grained, distinct, and not collapsed into broad Grounded Theory-style macro buckets?",
    },
    {
        "name": "Traceability and quote evidence",
        "human_prompt": "Can claims be traced to segment IDs and clean source quotes?",
    },
    {
        "name": "Memo-to-synthesis coherence",
        "human_prompt": "Do memos and final synthesis preserve tensions, variation, category relationships, and evidence-bounded claims?",
    },
    {
        "name": "Cross-coder divergence integration",
        "human_prompt": "If two coders were used, does integration explain convergence, divergence, and how disagreement changes category boundaries?",
    },
]

PROCESS_CRITERION_NAMES = {
    "Constant comparison process",
    "Boundary and negative-case handling",
    "Category differentiation and non-overmerge",
}
SYNTHESIS_CRITERION_NAMES = {
    "Memo-to-synthesis coherence",
    "Cross-coder divergence integration",
}

LLM_CRITERION_GROUPS = [
    {
        "name": "process_quality",
        "label": "CCA process quality",
        "criteria": [criterion for criterion in CRITERIA if criterion["name"] in PROCESS_CRITERION_NAMES],
    },
    {
        "name": "synthesis_integration_quality",
        "label": "Synthesis and integration quality",
        "criteria": [criterion for criterion in CRITERIA if criterion["name"] in SYNTHESIS_CRITERION_NAMES],
    },
]

FILE_REQUIREMENTS = [
    {
        "file_pattern": "analysis_summary.txt",
        "required": "Required",
        "used_for": "Run settings, coder count, high-level output counts.",
        "criteria": "Deterministic output inventory; Constant comparison process",
    },
    {
        "file_pattern": "segments_*.txt",
        "required": "Required for human review",
        "used_for": "Transcript context for segment IDs, quote checking, and human spot checks.",
        "criteria": "Traceability and quote evidence; Boundary and negative-case handling; Deterministic output inventory",
    },
    {
        "file_pattern": "incident_coding_coder*.txt",
        "required": "Required",
        "used_for": "Incident labels, comparison notes, analytic memos, boundary/negative incident flags.",
        "criteria": "Constant comparison process; Memo-to-synthesis coherence",
    },
    {
        "file_pattern": "incident_reconciliation_coder*.txt",
        "required": "Recommended",
        "used_for": "Global comparison pass, label clusters, cross-transcript comparison audit notes.",
        "criteria": "Constant comparison process; Deterministic output inventory",
    },
    {
        "file_pattern": "category_comparisons_coder*.txt",
        "required": "Required",
        "used_for": "Per-coder categories, supporting segments, quote evidence, boundary cases.",
        "criteria": "Category differentiation and non-overmerge; Boundary and negative-case handling; Traceability and quote evidence",
    },
    {
        "file_pattern": "cca_memos_coder*.txt",
        "required": "Required",
        "used_for": "Comparative memo quality, unresolved tensions, theoretical sampling leads.",
        "criteria": "Memo-to-synthesis coherence; Constant comparison process",
    },
    {
        "file_pattern": "cca_synthesis_coder*.txt",
        "required": "Required",
        "used_for": "Per-coder final narrative coherence and relationship claims.",
        "criteria": "Memo-to-synthesis coherence",
    },
    {
        "file_pattern": "integrated_incident_patterns.txt",
        "required": "Required for two-coder runs",
        "used_for": "Cross-coder incident pattern integration and convergence/divergence checks.",
        "criteria": "Cross-coder divergence integration; Constant comparison process",
    },
    {
        "file_pattern": "integrated_categories.txt",
        "required": "Required for two-coder runs",
        "used_for": "Integrated category granularity, distinctiveness, boundary/negative cases, divergence impact.",
        "criteria": "Cross-coder divergence integration; Category differentiation and non-overmerge; Boundary and negative-case handling",
    },
    {
        "file_pattern": "integrated_memo_digest.txt",
        "required": "Required for two-coder runs",
        "used_for": "Cross-coder memo synthesis and unresolved analytic tensions.",
        "criteria": "Cross-coder divergence integration; Memo-to-synthesis coherence",
    },
    {
        "file_pattern": "integrated_cca_summary.txt",
        "required": "Required for two-coder runs",
        "used_for": "Final integrated CCA narrative and evidence-bounded synthesis claims.",
        "criteria": "Memo-to-synthesis coherence; Cross-coder divergence integration",
    },
]

HUMAN_REVIEW_PROTOCOL = [
    {
        "criterion": "Constant comparison process",
        "files": "incident_reconciliation_coder*.txt; incident_coding_coder*.txt. Inspect 10 incident notes across early/middle/late transcripts plus all reconciliation summaries.",
        "review_task": "Check whether incident notes compare cases across local and non-local incidents and whether reconciliation improves comparison without erasing variation.",
        "minimum_review": "Inspect 10 incident notes across early/middle/late transcripts plus all reconciliation summaries.",
        "questions": [
            "Do comparison_notes explicitly name similarities and differences?",
            "Do later notes compare with non-adjacent or earlier transcript incidents?",
            "Are label clusters meaningful and not just superficial synonyms?",
            "Does reconciliation preserve minority, divergent, or boundary cases?",
        ],
    },
    {
        "criterion": "Traceability and quote evidence",
        "files": "category_comparisons_coder*.txt; integrated_categories.txt; segments_*.txt. Check at least 10 quote evidence items and their source segments.",
        "review_task": "Verify that category claims can be traced to segment IDs and clean source quotes.",
        "minimum_review": "Check at least 10 quote evidence items and their source segments.",
        "questions": [
            "Is each quote clean and free of speaker labels, line numbers, demographic headers, and question fragments?",
            "Does the quote accurately represent the surrounding segment context?",
            "Are important category claims supported by traceable segment IDs?",
        ],
    },
    {
        "criterion": "Boundary and negative-case handling",
        "files": "category_comparisons_coder*.txt; integrated_categories.txt; segments_*.txt. Review boundary_or_negative_cases for every final category.",
        "review_task": "Assess whether boundary/negative cases are present and analytically used, not merely listed.",
        "minimum_review": "Review boundary_or_negative_cases for every final category.",
        "questions": [
            "Does every final category include a boundary/negative case or a clear no-boundary reason?",
            "Do boundary cases change, narrow, split, or complicate the category definition?",
            "Are negative cases grounded in real segment context?",
        ],
    },
    {
        "criterion": "Category differentiation and non-overmerge",
        "files": "category_comparisons_coder*.txt; integrated_categories.txt; integrated_incident_patterns.txt. Review all final integrated categories if there are 7 or fewer; otherwise review at least 7 categories.",
        "review_task": "Judge whether final categories are mid-grained and distinct rather than broad Grounded Theory-style macro buckets.",
        "minimum_review": "Review all final integrated categories if there are 7 or fewer; otherwise review at least 7 categories.",
        "questions": [
            "Are distinct mechanisms preserved as separate categories when evidence supports separation?",
            "Are any categories redundant or too thin?",
            "Does the integrated category count fit the complexity of the evidence?",
        ],
    },
    {
        "criterion": "Memo-to-synthesis coherence",
        "files": "cca_memos_coder*.txt; integrated_memo_digest.txt; cca_synthesis_coder*.txt; integrated_cca_summary.txt. Read all final synthesis files, integrated memo digest, and at least 3 memos per coder when available.",
        "review_task": "Assess whether memos develop comparative insight and whether final synthesis follows from comparisons, categories, memos, and boundary cases without overclaiming.",
        "minimum_review": "Read all final synthesis files, integrated memo digest, and at least 3 memos per coder when available.",
        "questions": [
            "Do memos develop ideas beyond listing codes?",
            "Do memos preserve disagreement, variation, or unanswered questions?",
            "Does the synthesis explain relationships among categories rather than only list themes?",
            "Are synthesis claims limited to what the data and categories support?",
        ],
    },
    {
        "criterion": "Cross-coder divergence integration",
        "files": "integrated_incident_patterns.txt; integrated_categories.txt; integrated_memo_digest.txt; per-coder category files. Compare all integrated categories against both coders' category files.",
        "review_task": "For two-coder runs, check whether integration explains convergence, divergence, and category boundary impact.",
        "minimum_review": "Compare all integrated categories against both coders' category files.",
        "questions": [
            "Does integration preserve meaningful coder disagreement?",
            "Does divergence_boundary_impact explain whether disagreement narrows, splits, or complicates a category?",
            "Are integrated categories derived from both coders rather than one dominant coder?",
        ],
    },
    {
        "criterion": "Optional human theme/category alignment",
        "files": "integrated_categories.txt; integrated_cca_summary.txt; category_comparisons_coder*.txt; segments_*.txt; optional human codebook/theme checklist if the researcher has one",
        "review_task": (
            "If the researcher wants a human-reference comparison, the human reviewer creates or supplies reference themes/categories, "
            "then compares machine categories against them. The machine does not extract or score this reference automatically."
        ),
        "minimum_review": "Create 5-10 human reference themes/categories from human analysis or review all existing human codebook themes.",
        "questions": [
            "For each human theme, is the machine output strong, partial, or missing?",
            "Are any machine categories extra, weakly supported, or unsupported by transcript evidence?",
            "Does a different machine organization remain defensible as CCA, even if it does not match the human theme structure exactly?",
        ],
    },
]

LEGACY_EVALUATION_FILE_RE = re.compile(
    r"(?:cca_evaluation_coder\d+|integrated_evaluation|human_theme_checklist|human_reference_extraction|human_reference_findings)\.txt"
)


def _clip(value: Any, limit: int) -> str:
    text = str(value or "").strip()
    if len(text) <= limit:
        return text
    return text[:limit].rsplit(" ", 1)[0].rstrip()


def _clip_nested(value: Any, *, string_limit: int = 700, list_limit: int = 12, depth: int = 4) -> Any:
    if depth <= 0:
        return _clip(value, string_limit)
    if isinstance(value, dict):
        return {
            str(key): _clip_nested(item, string_limit=string_limit, list_limit=list_limit, depth=depth - 1)
            for key, item in value.items()
        }
    if isinstance(value, list):
        return [
            _clip_nested(item, string_limit=string_limit, list_limit=list_limit, depth=depth - 1)
            for item in value[:list_limit]
        ]
    if isinstance(value, str):
        return _clip(value, string_limit)
    return value


def _read_json(path: Path) -> Dict[str, Any]:
    try:
        return json.loads(path.read_text(encoding="utf-8", errors="replace"))
    except Exception:
        return {}


def _iter_output_files(input_dir: Path) -> Iterable[Path]:
    for path in sorted(input_dir.rglob("*")):
        if not path.is_file() or path.suffix.lower() not in {".txt", ".json"}:
            continue
        if LEGACY_EVALUATION_FILE_RE.fullmatch(path.name):
            continue
        if path.name in {"machine_evaluation.json"}:
            continue
        yield path


def _ignored_legacy_files(input_dir: Path) -> List[str]:
    ignored: List[str] = []
    for path in sorted(input_dir.rglob("*")):
        if path.is_file() and LEGACY_EVALUATION_FILE_RE.fullmatch(path.name):
            ignored.append(path.name)
    return sorted(set(ignored))


def _basename_map(input_dir: Path) -> Dict[str, Path]:
    out: Dict[str, Path] = {}
    for path in _iter_output_files(input_dir):
        out.setdefault(path.name, path)
    return out


def _find_named(files: Dict[str, Path], pattern: str) -> List[Tuple[str, Path]]:
    regex = re.compile(pattern)
    matched: List[Tuple[str, Path]] = []
    for name, path in files.items():
        m = regex.fullmatch(name)
        if m:
            matched.append((m.group(1) if m.groups() else name, path))
    return sorted(matched, key=lambda item: item[0])


def _category_records(data: Dict[str, Any]) -> List[Dict[str, Any]]:
    value = data.get("comparative_categories")
    if isinstance(value, list):
        return value
    value = data.get("categories")
    if isinstance(value, list):
        return value
    return []


def _normalize_for_match(value: Any) -> str:
    text = str(value or "").lower()
    text = re.sub(r"[^a-z0-9]+", " ", text)
    return re.sub(r"\s+", " ", text).strip()


def _quote_source_match(quote: str, source_text: str, threshold: float = 0.95) -> bool:
    quote_norm = _normalize_for_match(quote)
    source_norm = _normalize_for_match(source_text)
    if not quote_norm or not source_norm:
        return False
    if quote_norm in source_norm:
        return True
    match = SequenceMatcher(None, quote_norm, source_norm).find_longest_match(0, len(quote_norm), 0, len(source_norm))
    return (match.size / max(1, len(quote_norm))) >= threshold


def _quote_has_artifact(quote: str) -> bool:
    text = str(quote or "").strip()
    if not text:
        return True
    if re.match(r"^(?:participant|interviewer|moderator|respondent|patient|provider|nurse|doctor|researcher|p|i|r)\s*[:：-]", text, re.I):
        return True
    if re.match(r"^(?:line|ln\.?)?\s*\d{1,5}\s*[:.)-]", text, re.I):
        return True
    if re.match(r"^(?:age|gender|sex|race|ethnicity|education|income|site|location)\s*[:：-]", text, re.I):
        return True
    if text.endswith("?"):
        return True
    return False


def _count_quote_evidence(
    categories: List[Dict[str, Any]],
    segment_map: Optional[Dict[Tuple[str, int], str]] = None,
) -> Dict[str, Any]:
    quote_items = 0
    traceable_items = 0
    quote_strings = 0
    source_match_items = 0
    clean_quote_items = 0
    quote_audit_sample: List[Dict[str, Any]] = []
    boundary_categories = 0
    no_boundary_reason = 0
    for cat in categories:
        evidence = cat.get("supporting_quote_evidence") or []
        quotes = cat.get("supporting_quotes") or []
        if isinstance(quotes, list):
            quote_strings += len([q for q in quotes if str(q).strip()])
        if isinstance(evidence, list):
            quote_items += len(evidence)
            for item in evidence:
                if not isinstance(item, dict):
                    continue
                has_segment = bool(item.get("segment_id")) or (
                    bool(item.get("transcript")) and item.get("segment_number") is not None
                )
                quote = str(item.get("quote") or "").strip()
                source_context = str(item.get("source_context") or "").strip()
                segment_key = _parse_segment_ref(item)
                segment_text = segment_map.get(segment_key, "") if segment_map and segment_key else ""
                source_match = any(
                    _quote_source_match(quote, source)
                    for source in (source_context, segment_text)
                    if str(source or "").strip()
                )
                clean_quote = bool(quote) and not _quote_has_artifact(quote)
                if has_segment and quote and source_context:
                    traceable_items += 1
                if source_match:
                    source_match_items += 1
                if clean_quote:
                    clean_quote_items += 1
                if len(quote_audit_sample) < 20:
                    quote_audit_sample.append(
                        {
                            "category": cat.get("name"),
                            "segment_id": item.get("segment_id")
                            or (f"{item.get('transcript')}#{item.get('segment_number')}" if item.get("transcript") else ""),
                            "quote": _clip(quote, 220),
                            "has_segment_ref": has_segment,
                            "has_source_context": bool(source_context),
                            "source_match_95": source_match,
                            "clean_quote": clean_quote,
                        }
                    )
        boundaries = cat.get("boundary_or_negative_cases") or []
        if isinstance(boundaries, list) and boundaries:
            boundary_categories += 1
        elif str(cat.get("no_boundary_case_reason") or "").strip():
            no_boundary_reason += 1
    return {
        "quote_evidence_items": quote_items,
        "traceable_quote_evidence_items": traceable_items,
        "supporting_quote_strings": quote_strings,
        "quote_source_match_items": source_match_items,
        "clean_quote_items": clean_quote_items,
        "quote_audit_sample": quote_audit_sample,
        "categories_with_boundary_cases": boundary_categories,
        "categories_with_no_boundary_reason": no_boundary_reason,
    }


def _compact_category(cat: Dict[str, Any]) -> Dict[str, Any]:
    quote_metrics = _count_quote_evidence([cat])
    return {
        "name": cat.get("name"),
        "defining_properties": (cat.get("defining_properties") or cat.get("combined_properties") or [])[:6],
        "comparative_insights": (cat.get("comparative_insights") or [])[:6],
        "supporting_segments": (cat.get("supporting_segments") or [])[:6],
        "quote_evidence_items": quote_metrics["quote_evidence_items"],
        "traceable_quote_evidence_items": quote_metrics["traceable_quote_evidence_items"],
        "boundary_case_count": len(cat.get("boundary_or_negative_cases") or []),
        "no_boundary_case_reason": _clip(cat.get("no_boundary_case_reason", ""), 220),
        "coder_convergence": _clip(cat.get("coder_convergence", ""), 220),
        "coder_divergence": _clip(cat.get("coder_divergence", ""), 220),
        "divergence_boundary_impact": _clip(cat.get("divergence_boundary_impact", ""), 220),
    }


def _compact_incident(note: Dict[str, Any]) -> Dict[str, Any]:
    return {
        "transcript": note.get("transcript"),
        "segment_number": note.get("segment_number"),
        "labels": (note.get("labels") or [])[:8],
        "comparison_note_count": len(note.get("comparison_notes") or []),
        "has_global_comparison": bool(note.get("global_comparison_summary")),
        "has_boundary_or_negative_case": bool(note.get("boundary_or_negative_case")),
        "memo": _clip(note.get("analytic_memo", ""), 260),
    }


def _sample_list(items: List[Any], limit: int) -> List[Any]:
    if len(items) <= limit:
        return items
    if limit <= 1:
        return items[:1]
    step = (len(items) - 1) / (limit - 1)
    indexes = []
    for idx in range(limit):
        candidate = round(idx * step)
        if not indexes or candidate != indexes[-1]:
            indexes.append(candidate)
    while len(indexes) < limit:
        for candidate in range(len(items)):
            if candidate not in indexes:
                indexes.append(candidate)
                break
    return [items[idx] for idx in sorted(indexes[:limit])]


def _sanitize_analysis_summary(data: Dict[str, Any]) -> Dict[str, Any]:
    summary = str(data.get("summary") or "")
    if not summary:
        return data
    summary = re.sub(r"; reflective_evaluation = (?:true|false); human_reference = (?:true|false)", "", summary)
    summary = re.sub(r", evaluation=\d+", "", summary)
    summary = re.sub(r"; evaluation = \d+; human_theme_checklist = \d+", "", summary)
    clone = dict(data)
    clone["summary"] = summary
    return clone


def _segment_key(transcript: Any, segment_number: Any) -> Optional[Tuple[str, int]]:
    transcript_name = str(transcript or "").strip()
    if not transcript_name:
        return None
    try:
        seg_num = int(segment_number)
    except Exception:
        return None
    if seg_num <= 0:
        return None
    return (transcript_name, seg_num)


def _parse_segment_ref(value: Any) -> Optional[Tuple[str, int]]:
    if isinstance(value, dict):
        key = _segment_key(value.get("transcript"), value.get("segment_number"))
        if key:
            return key
        segment_id = str(value.get("segment_id") or "").strip()
        if segment_id:
            return _parse_segment_ref(segment_id)
        return None
    text = str(value or "").strip()
    if not text:
        return None
    match = re.match(r"^(.+?)[#:]([0-9]+)$", text)
    if not match:
        return None
    return _segment_key(match.group(1).strip(), match.group(2))


def _load_segment_context(files: Dict[str, Path]) -> Tuple[Dict[Tuple[str, int], str], List[Dict[str, Any]]]:
    segment_map: Dict[Tuple[str, int], str] = {}
    summaries: List[Dict[str, Any]] = []
    for name, path in sorted(files.items()):
        if not (name.startswith("segments_") and name.endswith(".txt")):
            continue
        data = _read_json(path)
        transcript = str(data.get("transcript") or name.removeprefix("segments_").removesuffix(".txt")).strip()
        segments = data.get("segments") or {}
        if not isinstance(segments, dict):
            continue
        loaded = 0
        for raw_num, raw_text in segments.items():
            key = _segment_key(transcript, raw_num)
            if not key:
                continue
            segment_map[key] = str(raw_text or "")
            loaded += 1
        summaries.append({"file": name, "transcript": transcript, "segments": loaded})
    return segment_map, summaries


def _add_segment_ref(refs: set[Tuple[str, int]], value: Any) -> None:
    key = _parse_segment_ref(value)
    if key:
        refs.add(key)


def _collect_referenced_segments(
    *,
    incidents: List[Dict[str, Any]],
    categories: List[Dict[str, Any]],
    integrated_categories: List[Dict[str, Any]],
    incident_patterns: List[Dict[str, Any]],
) -> set[Tuple[str, int]]:
    refs: set[Tuple[str, int]] = set()
    for note in incidents:
        _add_segment_ref(refs, note)
    for cat in [*categories, *integrated_categories]:
        for field in ("supporting_segments", "boundary_or_negative_cases", "supporting_quote_evidence"):
            for item in cat.get(field) or []:
                _add_segment_ref(refs, item)
        for item in cat.get("supporting_quote_evidence") or []:
            if isinstance(item, dict):
                _add_segment_ref(refs, item.get("segment_id"))
    for pattern in incident_patterns:
        for item in pattern.get("representative_segments") or []:
            _add_segment_ref(refs, item)
    return refs


def _build_segment_context_preview(
    *,
    refs: set[Tuple[str, int]],
    segment_map: Dict[Tuple[str, int], str],
    limit: int = 30,
) -> Tuple[List[Dict[str, Any]], List[str]]:
    preview: List[Dict[str, Any]] = []
    missing: List[str] = []
    sampled_refs = _sample_list(sorted(refs), limit)
    for transcript, segment_number in sorted(refs):
        text = segment_map.get((transcript, segment_number))
        ref = f"{transcript}#{segment_number}"
        if text is None:
            missing.append(ref)
    for transcript, segment_number in sampled_refs:
        text = segment_map.get((transcript, segment_number))
        if text is None:
            continue
        ref = f"{transcript}#{segment_number}"
        preview.append(
            {
                "segment_id": ref,
                "transcript": transcript,
                "segment_number": segment_number,
                "text": _clip(text, 700),
            }
        )
    return preview, missing


def load_cca_run(input_dir: Path) -> Dict[str, Any]:
    files = _basename_map(input_dir)
    ignored_legacy_files = _ignored_legacy_files(input_dir)
    coder_ids = sorted(
        {
            coder
            for coder, _path in _find_named(files, r"incident_coding_(coder\d+)\.txt")
        }
    )
    coder_outputs: Dict[str, Dict[str, Any]] = {}
    all_incidents: List[Dict[str, Any]] = []
    all_categories: List[Dict[str, Any]] = []
    all_memos: List[Dict[str, Any]] = []
    all_reconciliations: List[Dict[str, Any]] = []

    for coder in coder_ids:
        incident_data = _read_json(files.get(f"incident_coding_{coder}.txt", Path()))
        category_data = _read_json(files.get(f"category_comparisons_{coder}.txt", Path()))
        memo_data = _read_json(files.get(f"cca_memos_{coder}.txt", Path()))
        synthesis_data = _read_json(files.get(f"cca_synthesis_{coder}.txt", Path()))
        reconciliation_data = _read_json(files.get(f"incident_reconciliation_{coder}.txt", Path()))

        incidents = incident_data.get("incident_notes") or []
        categories = _category_records(category_data)
        memos = memo_data.get("comparative_memos") or []
        all_incidents.extend(incidents)
        all_categories.extend(categories)
        all_memos.extend(memos)
        if reconciliation_data:
            all_reconciliations.append(reconciliation_data)
        coder_outputs[coder] = {
            "incident_notes": incidents,
            "categories": categories,
            "memos": memos,
            "synthesis": synthesis_data,
            "reconciliation": reconciliation_data,
        }

    integrated_categories = _category_records(_read_json(files.get("integrated_categories.txt", Path())))
    integrated_patterns = (_read_json(files.get("integrated_incident_patterns.txt", Path())).get("incident_patterns") or [])
    integrated_memos = (_read_json(files.get("integrated_memo_digest.txt", Path())).get("memo_digest") or [])
    integrated_summary = _read_json(files.get("integrated_cca_summary.txt", Path()))
    analysis_summary = _sanitize_analysis_summary(_read_json(files.get("analysis_summary.txt", Path())))
    segment_map, segment_file_summaries = _load_segment_context(files)

    integrated_present = bool(integrated_categories or integrated_patterns or integrated_summary)
    required_files = ["analysis_summary.txt"]
    for coder in coder_ids:
        required_files.extend(
            [
                f"incident_coding_{coder}.txt",
                f"category_comparisons_{coder}.txt",
                f"cca_memos_{coder}.txt",
                f"cca_synthesis_{coder}.txt",
            ]
        )
    if len(coder_ids) > 1:
        required_files.extend(
            [
                "integrated_incident_patterns.txt",
                "integrated_categories.txt",
                "integrated_memo_digest.txt",
                "integrated_cca_summary.txt",
            ]
        )
    if segment_file_summaries:
        required_files.extend(summary["file"] for summary in segment_file_summaries)
    else:
        required_files.append("segments_*.txt (at least one transcript segment file)")
    missing_files = [name for name in required_files if name not in files]
    if segment_file_summaries:
        missing_files = [name for name in missing_files if not name.startswith("segments_")]

    referenced_segments = _collect_referenced_segments(
        incidents=all_incidents,
        categories=all_categories,
        integrated_categories=integrated_categories,
        incident_patterns=integrated_patterns,
    )
    segment_context_preview, missing_referenced_segments = _build_segment_context_preview(
        refs=referenced_segments,
        segment_map=segment_map,
    )

    quote_metrics = _count_quote_evidence(all_categories, segment_map=segment_map)
    integrated_boundary_metrics = _count_quote_evidence(integrated_categories, segment_map=segment_map)
    incident_label_count = sum(len(note.get("labels") or []) for note in all_incidents)
    incident_comparison_count = sum(len(note.get("comparison_notes") or []) for note in all_incidents)
    substantive_comparison_count = 0
    for note in all_incidents:
        for comparison in note.get("comparison_notes") or []:
            if not isinstance(comparison, dict):
                continue
            similarities = str(comparison.get("similarities") or "").strip().lower()
            differences = str(comparison.get("differences") or "").strip().lower()
            focus = str(comparison.get("focus") or "").strip().lower()
            if "no prior" in focus or similarities.startswith("n/a"):
                continue
            if similarities or differences:
                substantive_comparison_count += 1
    global_comparison_count = sum(1 for note in all_incidents if note.get("global_comparison_summary"))
    incident_boundary_count = sum(1 for note in all_incidents if note.get("boundary_or_negative_case"))
    reconciliation_label_clusters = sum(len(item.get("label_clusters") or []) for item in all_reconciliations)
    reconciliation_global_rows = sum(len(item.get("global_comparisons") or []) for item in all_reconciliations)
    integrated_divergence_boundary_count = sum(
        1 for cat in integrated_categories if str(cat.get("divergence_boundary_impact") or "").strip()
    )
    integrated_boundary_or_reason_count = (
        integrated_boundary_metrics.get("categories_with_boundary_cases", 0)
        + integrated_boundary_metrics.get("categories_with_no_boundary_reason", 0)
    )

    metrics = {
        "input_dir": str(input_dir),
        "file_inventory": {
            "total_output_files": len(files),
            "present_files": sorted(files.keys()),
            "ignored_legacy_evaluation_files": ignored_legacy_files,
            "required_files": required_files,
            "missing_required_files": missing_files,
        },
        "counts": {
            "coders": len(coder_ids),
            "incident_notes": len(all_incidents),
            "incident_labels": incident_label_count,
            "incident_comparison_notes": incident_comparison_count,
            "substantive_incident_comparison_notes": substantive_comparison_count,
            "incident_global_comparisons": global_comparison_count,
            "incident_boundary_or_negative_cases": incident_boundary_count,
            "reconciliation_files": len(all_reconciliations),
            "reconciliation_label_clusters": reconciliation_label_clusters,
            "reconciliation_global_comparisons": reconciliation_global_rows,
            "per_coder_categories": len(all_categories),
            "per_coder_memos": len(all_memos),
            "integrated_present": integrated_present,
            "integrated_incident_patterns": len(integrated_patterns),
            "integrated_categories": len(integrated_categories),
            "integrated_memos": len(integrated_memos),
            "has_integrated_summary": bool(str(integrated_summary.get("comparative_summary") or "").strip()),
            "segment_files": len(segment_file_summaries),
            "segments_loaded": len(segment_map),
            "referenced_segments": len(referenced_segments),
            "missing_referenced_segments": len(missing_referenced_segments),
            "integrated_categories_with_boundary_or_reason": integrated_boundary_or_reason_count,
            "integrated_categories_with_divergence_boundary_impact": integrated_divergence_boundary_count,
        },
        "quote_traceability": quote_metrics,
        "integrated_boundary_metrics": integrated_boundary_metrics,
    }

    compact = {
        "metrics": metrics,
        "analysis_summary": analysis_summary,
        "segment_context": {
            "segment_files": segment_file_summaries,
            "referenced_segment_preview": segment_context_preview,
            "missing_referenced_segments": missing_referenced_segments[:80],
        },
        "coder_outputs": {
            coder: {
                "incident_notes_sample": [_compact_incident(note) for note in _sample_list(data["incident_notes"], 24)],
                "categories": [_compact_category(cat) for cat in data["categories"][:14]],
                "memos": data["memos"][:8],
                "synthesis": data["synthesis"],
                "reconciliation_summary": {
                    "label_clusters": (data["reconciliation"].get("label_clusters") or [])[:20],
                    "audit_note": data["reconciliation"].get("audit_note", ""),
                },
            }
            for coder, data in coder_outputs.items()
        },
        "integrated_outputs": {
            "incident_patterns": integrated_patterns[:40],
            "categories": [_compact_category(cat) for cat in integrated_categories[:14]],
            "memo_digest": integrated_memos[:10],
            "summary": integrated_summary,
        },
    }
    return compact


def _deterministic_score(metrics: Dict[str, Any]) -> Dict[str, Any]:
    counts = metrics.get("counts", {})
    inventory = metrics.get("file_inventory", {})
    quotes = metrics.get("quote_traceability", {})
    boundaries = metrics.get("integrated_boundary_metrics", {})

    missing = len(inventory.get("missing_required_files") or [])
    incident_notes = int(counts.get("incident_notes") or 0)
    comparison_notes = int(counts.get("incident_comparison_notes") or 0)
    substantive_comparisons = int(counts.get("substantive_incident_comparison_notes") or 0)
    categories = int(counts.get("per_coder_categories") or 0)
    traceable_quotes = int(quotes.get("traceable_quote_evidence_items") or 0)
    quote_items = int(quotes.get("quote_evidence_items") or 0)
    source_matches = int(quotes.get("quote_source_match_items") or 0)
    clean_quotes = int(quotes.get("clean_quote_items") or 0)
    integrated_categories = int(counts.get("integrated_categories") or 0)
    boundary_categories = int(boundaries.get("categories_with_boundary_cases") or 0)
    boundary_or_reason = int(counts.get("integrated_categories_with_boundary_or_reason") or 0)
    divergence_boundary = int(counts.get("integrated_categories_with_divergence_boundary_impact") or 0)
    reconciliation_files = int(counts.get("reconciliation_files") or 0)
    reconciliation_clusters = int(counts.get("reconciliation_label_clusters") or 0)
    reconciliation_rows = int(counts.get("reconciliation_global_comparisons") or 0)
    segment_files = int(counts.get("segment_files") or 0)
    missing_segment_refs = int(counts.get("missing_referenced_segments") or 0)

    incident_comparison = 5 if incident_notes and substantive_comparisons >= incident_notes * 0.75 else 4 if substantive_comparisons else 2
    memory_coverage = 5 if substantive_comparisons and reconciliation_rows else 4 if substantive_comparisons else 2
    reconciliation = 5 if reconciliation_files and reconciliation_clusters and reconciliation_rows else 4 if reconciliation_files and reconciliation_clusters else 2
    comparison_process = round((incident_comparison + memory_coverage + reconciliation) / 3)
    if quote_items:
        traceable_rate = traceable_quotes / quote_items
        source_match_rate = source_matches / quote_items
        clean_rate = clean_quotes / quote_items
        if traceable_rate >= 0.95 and source_match_rate >= 0.95 and clean_rate >= 0.90:
            traceability = 5
        elif traceable_rate >= 0.80 and source_match_rate >= 0.80 and clean_rate >= 0.75:
            traceability = 4
        elif traceable_quotes or source_matches:
            traceability = 3
        else:
            traceability = 2
    else:
        traceability = 2
    boundary = 5 if integrated_categories and boundary_or_reason >= integrated_categories else 4 if boundary_categories else 2
    differentiation = 5 if 5 <= integrated_categories <= 7 else 4 if categories and integrated_categories else 3 if categories else 1
    memo_synthesis = 4 if counts.get("per_coder_memos") and (counts.get("has_integrated_summary") or counts.get("coders") == 1) else 3
    if counts.get("has_integrated_summary") and counts.get("per_coder_memos"):
        memo_synthesis = 5 if boundary_or_reason else 4
    divergence = 5 if counts.get("coders", 0) <= 1 else 5 if integrated_categories and divergence_boundary >= integrated_categories else 3 if integrated_categories else 2
    completeness = 5 if missing == 0 else 3 if missing <= 2 else 1
    if segment_files == 0:
        completeness = min(completeness, 3)
        traceability = min(traceability, 3)
    if missing_segment_refs:
        traceability = min(traceability, 4)

    rows = [
        ("Constant comparison process", comparison_process),
        ("Boundary and negative-case handling", boundary),
        ("Category differentiation and non-overmerge", differentiation),
        ("Traceability and quote evidence", traceability),
        ("Memo-to-synthesis coherence", memo_synthesis),
        ("Cross-coder divergence integration", divergence),
    ]
    criteria = []
    for name, score in rows:
        criterion = {
            "name": name,
            "score": score,
            "rationale": "Deterministic fallback score based on output counts and required-file checks.",
            "evidence": [],
            "risks": ["No LLM narrative evaluation was run."],
            "recommendations": ["Run the evaluation app with an OpenAI API key for a fuller machine review."],
        }
        if name == "Traceability and quote evidence":
            criterion.update(
                {
                    "rationale": (
                        "Deterministic quote audit: checks segment/source-context fields, loose 95% quote-to-source matching, "
                        "and common transcript artifacts in quote text."
                    ),
                    "evidence": [
                        f"{traceable_quotes}/{quote_items} quote evidence items include segment/source context.",
                        f"{source_matches}/{quote_items} quote evidence items match source context or segment text at the 95% loose-match threshold.",
                        f"{clean_quotes}/{quote_items} quote evidence items pass the clean-quote artifact check.",
                    ],
                    "risks": [
                        "This code check cannot judge whether the quote is the best evidence for the category claim.",
                        "No LLM narrative evaluation was run.",
                    ],
                }
            )
        criteria.append(criterion)
    overall = round(sum(score for _name, score in rows) / len(rows), 2)
    return {
        "overall_score": overall,
        "machine_summary": "Deterministic fallback evaluation only. Treat as a structural completeness check, not a qualitative judgment.",
        "criteria": criteria,
        "strengths": [],
        "limitations": ["No LLM-based qualitative audit was run."],
        "human_review_priorities": ["Complete the blank human review section in the HTML report."],
    }


def _criterion_by_name(evaluation: Dict[str, Any], name: str) -> Optional[Dict[str, Any]]:
    for criterion in evaluation.get("criteria") or []:
        if criterion.get("name") == name:
            return criterion
    return None


def _coerce_score(value: Any) -> int:
    try:
        score = int(round(float(value)))
    except Exception:
        return 1
    return max(1, min(5, score))


def _normalize_criterion_score(criterion: Dict[str, Any]) -> Dict[str, Any]:
    normalized = dict(criterion)
    normalized["score"] = _coerce_score(normalized.get("score"))
    return normalized


def _deterministic_traceability_criterion(metrics: Dict[str, Any]) -> Dict[str, Any]:
    deterministic = _deterministic_score(metrics)
    criterion = dict(_criterion_by_name(deterministic, "Traceability and quote evidence") or {})
    quotes = metrics.get("quote_traceability") or {}
    quote_items = int(quotes.get("quote_evidence_items") or 0)
    source_matches = int(quotes.get("quote_source_match_items") or 0)
    clean_quotes = int(quotes.get("clean_quote_items") or 0)
    traceable_quotes = int(quotes.get("traceable_quote_evidence_items") or 0)
    criterion["rationale"] = (
        "Deterministic quote audit: checks segment/source-context fields, loose 95% quote-to-source matching, "
        "and common transcript artifacts in quote text."
    )
    criterion["evidence"] = [
        f"{traceable_quotes}/{quote_items} quote evidence items include segment/source context.",
        f"{source_matches}/{quote_items} quote evidence items match source context or segment text at the 95% loose-match threshold.",
        f"{clean_quotes}/{quote_items} quote evidence items pass the clean-quote artifact check.",
    ]
    criterion["risks"] = [
        "This code check cannot judge whether the quote is the best evidence for the category claim."
    ]
    criterion["recommendations"] = [
        "Use the blank human review section to spot-check whether matched quotes substantively support the category interpretation."
    ]
    return criterion


def _merge_machine_and_deterministic_evaluation(
    *,
    llm_eval: Dict[str, Any],
    deterministic_eval: Dict[str, Any],
    traceability: Dict[str, Any],
) -> Dict[str, Any]:
    llm_by_name = {criterion.get("name"): criterion for criterion in llm_eval.get("criteria") or []}
    deterministic_by_name = {criterion.get("name"): criterion for criterion in deterministic_eval.get("criteria") or []}
    merged_criteria: List[Dict[str, Any]] = []
    for rubric in CRITERIA:
        name = rubric["name"]
        if name == "Traceability and quote evidence":
            merged_criteria.append(_normalize_criterion_score(traceability))
        else:
            merged_criteria.append(
                _normalize_criterion_score(
                    llm_by_name.get(name) or deterministic_by_name.get(name) or {"name": name, "score": 1}
                )
            )
    scores: List[float] = []
    for criterion in merged_criteria:
        try:
            scores.append(float(criterion.get("score")))
        except Exception:
            continue
    overall = round(sum(scores) / len(scores), 2) if scores else 0
    limitations = list(llm_eval.get("limitations") or [])
    limitations.append("Traceability and quote evidence is code-scored; LLM does not independently score that criterion.")
    return {
        "overall_score": overall,
        "machine_summary": llm_eval.get("machine_summary", ""),
        "criteria": merged_criteria,
        "strengths": llm_eval.get("strengths") or [],
        "limitations": limitations,
        "human_review_priorities": llm_eval.get("human_review_priorities") or [],
    }


def _criterion_group_context(compact_run: Dict[str, Any], group_name: str) -> Dict[str, Any]:
    metrics = compact_run.get("metrics") or {}
    counts = metrics.get("counts") or {}
    base = {
        "metrics": {
            "file_inventory": metrics.get("file_inventory") or {},
            "counts": counts,
            "quote_traceability": metrics.get("quote_traceability") or {},
            "integrated_boundary_metrics": metrics.get("integrated_boundary_metrics") or {},
        },
        "analysis_summary": compact_run.get("analysis_summary") or {},
    }
    if group_name == "process_quality":
        coder_packets: Dict[str, Any] = {}
        for coder, data in (compact_run.get("coder_outputs") or {}).items():
            coder_packets[coder] = {
                "incident_notes_sample": _sample_list(data.get("incident_notes_sample") or [], 12),
                "categories": (data.get("categories") or [])[:8],
                "reconciliation_summary": _clip_nested(data.get("reconciliation_summary") or {}, string_limit=450, list_limit=6),
            }
        return {
            **base,
            "segment_context": {
                "segment_files": (compact_run.get("segment_context") or {}).get("segment_files") or [],
                "referenced_segment_preview": _sample_list(
                    (compact_run.get("segment_context") or {}).get("referenced_segment_preview") or [],
                    12,
                ),
                "missing_referenced_segments": (compact_run.get("segment_context") or {}).get("missing_referenced_segments") or [],
            },
            "coder_outputs": coder_packets,
            "integrated_outputs": {
                "incident_patterns": _clip_nested(
                    (compact_run.get("integrated_outputs") or {}).get("incident_patterns") or [],
                    string_limit=450,
                    list_limit=8,
                ),
                "categories": ((compact_run.get("integrated_outputs") or {}).get("categories") or [])[:8],
            },
        }
    if group_name == "synthesis_integration_quality":
        coder_packets = {}
        for coder, data in (compact_run.get("coder_outputs") or {}).items():
            coder_packets[coder] = {
                "categories": (data.get("categories") or [])[:8],
                "memos": _clip_nested(data.get("memos") or [], string_limit=700, list_limit=6),
                "synthesis": _clip_nested(data.get("synthesis") or {}, string_limit=900, list_limit=10),
            }
        return {
            **base,
            "coder_outputs": coder_packets,
            "integrated_outputs": {
                "categories": ((compact_run.get("integrated_outputs") or {}).get("categories") or [])[:8],
                "memo_digest": _clip_nested(
                    (compact_run.get("integrated_outputs") or {}).get("memo_digest") or [],
                    string_limit=700,
                    list_limit=10,
                ),
                "summary": _clip_nested((compact_run.get("integrated_outputs") or {}).get("summary") or {}, string_limit=900, list_limit=10),
            },
        }
    return base


def _run_llm_criterion_group(
    *,
    sdk: AgentSDK,
    system: str,
    compact_run: Dict[str, Any],
    metrics: Dict[str, Any],
    group: Dict[str, Any],
) -> Dict[str, Any]:
    group_name = str(group.get("name") or "")
    criteria = group.get("criteria") or []
    payload = {
        "rubric_context": {
            "no_ground_truth": True,
            "criterion_group": group.get("label") or group_name,
            "criteria": criteria,
            "score_type": "integer only; allowed values are exactly 1, 2, 3, 4, or 5",
            "score_scale": SCORING_SCALE,
            "criterion_score_anchors": {
                criterion.get("name"): CRITERION_SCORE_ANCHORS.get(criterion.get("name"), {})
                for criterion in criteria
            },
            "code_scored_criteria": [
                {
                    "name": "Traceability and quote evidence",
                    "method": "Deterministic quote audit using segment/source-context presence, loose 95% quote-to-source matching, and common quote artifact checks.",
                    "metrics": metrics.get("quote_traceability") or {},
                }
            ],
        },
        "cca_output_packet": _criterion_group_context(compact_run, group_name),
        "instructions": (
            "Evaluate only the criteria supplied in rubric_context.criteria for this criterion group. "
            "Do not score Traceability and quote evidence; it is code-scored and will be merged into the final report. "
            "Do not score output completeness separately; it is checked deterministically in the file inventory. "
            "For each returned criterion, score must be an integer only: 1, 2, 3, 4, or 5. Do not use decimals, ranges, percentages, or text labels as scores. "
            "Apply both the general score_scale and the criterion-specific criterion_score_anchors; the criterion-specific anchors are authoritative when they are more precise. "
            "If evidence is mixed or uncertain, choose the lower defensible score and explain the uncertainty in risks. "
            "Use only the supplied criterion-specific packet, and cite concrete file names, category names, counts, or segment IDs. "
            "Return JSON only using the requested schema."
        ),
        "schema": MACHINE_EVALUATION_SCHEMA,
    }
    data = sdk.run_json(
        system=system,
        user=json.dumps(payload, ensure_ascii=False),
        schema_hint=MACHINE_EVALUATION_SCHEMA,
        attempts=1,
        timeout_s=220.0,
    )
    return data or {}


def run_machine_evaluation(
    *,
    compact_run: Dict[str, Any],
    api_key: Optional[str],
    model: str,
    use_llm: bool = True,
) -> Dict[str, Any]:
    metrics = compact_run.get("metrics") or {}
    deterministic_eval = _deterministic_score(metrics)
    if not use_llm or not api_key:
        return deterministic_eval

    sdk = AgentSDK(model=model, api_key=api_key)
    system = (
        "You are an independent qualitative-methods audit agent for a Constant Comparative Analysis output. "
        "The researcher has no ground truth and no published human answer. Evaluate only whether the uploaded CCA output "
        "is methodologically credible, traceable, coherent, and usable for human qualitative research review. "
        "This is not a Grounded Theory replication audit: do not reward outputs merely for producing broad themes or a core theory. "
        "Prioritize CCA-specific behavior: incident-to-incident comparison, iterative differentiation, negative/boundary cases, "
        "cross-case memory, global reconciliation, category boundary clarity, and divergence-aware integration. "
        "Do not compare against an article or invent missing ground truth. Machine scores are provisional and must leave room for human judgment. "
        "Score each criterion from 1 to 5. Return JSON only."
    )
    group_results = [
        _run_llm_criterion_group(
            sdk=sdk,
            system=system,
            compact_run=compact_run,
            metrics=metrics,
            group=group,
        )
        for group in LLM_CRITERION_GROUPS
    ]
    if not any(group_results):
        return deterministic_eval
    criteria: List[Dict[str, Any]] = []
    strengths: List[str] = []
    limitations: List[str] = []
    human_review_priorities: List[str] = []
    summaries: List[str] = []
    for result in group_results:
        criteria.extend(result.get("criteria") or [])
        strengths.extend(result.get("strengths") or [])
        limitations.extend(result.get("limitations") or [])
        human_review_priorities.extend(result.get("human_review_priorities") or [])
        if str(result.get("machine_summary") or "").strip():
            summaries.append(str(result.get("machine_summary")).strip())
    llm_eval = {
        "overall_score": 0,
        "machine_summary": "Two criterion-specific LLM calls were used: process quality and synthesis/integration quality. " + " ".join(summaries),
        "criteria": criteria,
        "strengths": strengths,
        "limitations": limitations,
        "human_review_priorities": human_review_priorities,
    }
    return _merge_machine_and_deterministic_evaluation(
        llm_eval=llm_eval,
        deterministic_eval=deterministic_eval,
        traceability=_deterministic_traceability_criterion(metrics),
    )


def _esc(value: Any) -> str:
    return html.escape(str(value if value is not None else ""))


def _score_class(score: Any) -> str:
    try:
        val = float(score)
    except Exception:
        return "score-low"
    if val >= 4:
        return "score-high"
    if val >= 3:
        return "score-mid"
    return "score-low"


def _human_review_items() -> List[Dict[str, Any]]:
    protocol_by_criterion = {item.get("criterion"): item for item in HUMAN_REVIEW_PROTOCOL}
    items: List[Dict[str, Any]] = []
    for idx, criterion in enumerate(CRITERIA, start=1):
        name = criterion["name"]
        protocol = protocol_by_criterion.get(name, {})
        items.append(
            {
                "number": idx,
                "criterion": name,
                "prompt": criterion.get("human_prompt", ""),
                "files": protocol.get("files", ""),
                "anchors": CRITERION_SCORE_ANCHORS.get(name, {}),
            }
        )
    return items


def _human_review_sections() -> List[Tuple[str, List[Dict[str, Any]]]]:
    items = {item["criterion"]: item for item in _human_review_items()}
    return [
        (
            "A. CCA Process Quality",
            [
                items["Constant comparison process"],
                items["Boundary and negative-case handling"],
                items["Category differentiation and non-overmerge"],
            ],
        ),
        ("B. Traceability and Quote Evidence", [items["Traceability and quote evidence"]]),
        (
            "C. Synthesis and Integration Quality",
            [
                items["Memo-to-synthesis coherence"],
                items["Cross-coder divergence integration"],
            ],
        ),
    ]


def write_human_evaluation_docx(path: Path, *, generated_at: str = "") -> Path:
    try:
        from docx import Document  # type: ignore
        from docx.shared import Pt, RGBColor  # type: ignore
    except Exception as exc:
        raise RuntimeError(f"python-docx is required to write the human evaluation form: {exc}") from exc

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Pt(54)
    section.bottom_margin = Pt(54)
    section.left_margin = Pt(63)
    section.right_margin = Pt(63)

    styles = doc.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(9.8)
    styles["Normal"].paragraph_format.space_after = Pt(3)
    styles["Normal"].paragraph_format.line_spacing = 1.08
    styles["Heading 1"].font.name = "Calibri"
    styles["Heading 1"].font.size = Pt(13)
    styles["Heading 1"].font.color.rgb = RGBColor(11, 37, 69)
    styles["Heading 1"].paragraph_format.space_before = Pt(8)
    styles["Heading 1"].paragraph_format.space_after = Pt(4)

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(7)
    run = title.add_run("CCA Human Evaluation Rubric (1-5 Likert Scale)")
    run.bold = True
    run.font.name = "Calibri"
    run.font.size = Pt(15)
    run.font.color.rgb = RGBColor(11, 37, 69)
    if generated_at:
        p = doc.add_paragraph(f"Generated: {generated_at}")
        p.runs[0].font.color.rgb = RGBColor(100, 116, 139)

    for section_title, items in _human_review_sections():
        doc.add_heading(section_title, level=1)
        for item in items:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            r = p.add_run(f"{item['number']}. {item['criterion']}\n")
            r.bold = True
            r.font.name = "Calibri"
            r.font.size = Pt(9.8)
            r.font.color.rgb = RGBColor(17, 24, 39)
            q = p.add_run(f"{item['prompt']}\n")
            q.font.name = "Calibri"
            q.font.size = Pt(9.8)
            f = p.add_run(f"Files to inspect: {item['files']}\n")
            f.font.name = "Calibri"
            f.font.size = Pt(9.6)
            f.italic = True
            a = p.add_run("Anchors:")
            a.font.name = "Calibri"
            a.font.size = Pt(9.8)
            for score, text in item["anchors"].items():
                ap = doc.add_paragraph(f"{score} = {text}")
                ap.paragraph_format.left_indent = Pt(14)
                ap.paragraph_format.space_after = Pt(0)
                ap.paragraph_format.line_spacing = 1.02
                for ar in ap.runs:
                    ar.font.name = "Calibri"
                    ar.font.size = Pt(9.2)

    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(path))
    return path


def _list_html(items: Iterable[Any]) -> str:
    values = [str(item).strip() for item in items if str(item).strip()]
    if not values:
        return "<p class=\"muted\">None reported.</p>"
    return "<ul>" + "".join(f"<li>{_esc(item)}</li>" for item in values) + "</ul>"


def _metric_card(label: str, value: Any, note: str = "") -> str:
    return (
        "<div class=\"metric-card\">"
        f"<div class=\"metric-label\">{_esc(label)}</div>"
        f"<div class=\"metric-value\">{_esc(value)}</div>"
        f"<div class=\"metric-note\">{_esc(note)}</div>"
        "</div>"
    )


def _fmt_count(value: Any) -> int:
    try:
        return int(value or 0)
    except Exception:
        return 0


def _supporting_metrics_for_domain(name: str, counts: Dict[str, Any], quotes: Dict[str, Any], boundaries: Dict[str, Any]) -> List[str]:
    """Return non-overlapping, paper-facing audit metrics for each evaluation domain."""
    if name == "Constant comparison process":
        return [
            f"{_fmt_count(counts.get('substantive_incident_comparison_notes'))} substantive comparison notes",
            f"{_fmt_count(counts.get('reconciliation_label_clusters'))} reconciliation clusters",
            f"{_fmt_count(counts.get('reconciliation_global_comparisons'))} global reconciliation comparisons",
        ]
    if name == "Traceability and quote evidence":
        quote_items = _fmt_count(quotes.get("quote_evidence_items"))
        return [
            f"{_fmt_count(quotes.get('traceable_quote_evidence_items'))}/{quote_items} quote evidence items traceable",
            f"{_fmt_count(quotes.get('quote_source_match_items'))}/{quote_items} quote-source matches",
            f"{_fmt_count(counts.get('missing_referenced_segments'))} missing referenced segments",
        ]
    if name == "Boundary and negative-case handling":
        integrated_categories = _fmt_count(counts.get("integrated_categories"))
        return [
            f"{_fmt_count(counts.get('integrated_categories_with_boundary_or_reason'))}/{integrated_categories} final categories with boundary case or rationale",
            f"{_fmt_count(boundaries.get('categories_with_boundary_cases'))} integrated boundary-case categories",
            f"{_fmt_count(boundaries.get('categories_with_no_boundary_reason'))} explicit no-boundary rationales",
        ]
    if name == "Category differentiation and non-overmerge":
        return [
            f"{_fmt_count(counts.get('per_coder_categories'))} per-coder categories",
            f"{_fmt_count(counts.get('integrated_categories'))} integrated categories",
        ]
    if name == "Memo-to-synthesis coherence":
        return [
            f"{_fmt_count(counts.get('per_coder_memos'))} comparative memos",
            f"{'Yes' if counts.get('has_integrated_summary') else 'No'} integrated summary present",
            f"{_fmt_count(counts.get('integrated_memos'))} integrated memo-digest entries",
        ]
    if name == "Cross-coder divergence integration":
        integrated_categories = _fmt_count(counts.get("integrated_categories"))
        return [
            f"{_fmt_count(counts.get('coders'))} coders",
            f"{_fmt_count(counts.get('integrated_categories_with_divergence_boundary_impact'))}/{integrated_categories} categories with divergence-boundary impact",
            f"{'Yes' if counts.get('integrated_present') else 'No'} integrated two-coder output present",
        ]
    return []


def _criterion_group_and_number(name: str) -> Tuple[str, int]:
    mapping = {
        "Constant comparison process": ("A. CCA Process Quality", 1),
        "Boundary and negative-case handling": ("A. CCA Process Quality", 2),
        "Category differentiation and non-overmerge": ("A. CCA Process Quality", 3),
        "Traceability and quote evidence": ("B. Traceability and Quote Evidence", 4),
        "Memo-to-synthesis coherence": ("C. Synthesis and Integration Quality", 5),
        "Cross-coder divergence integration": ("C. Synthesis and Integration Quality", 6),
    }
    return mapping.get(name, ("", 0))


def render_html_report(
    *,
    compact_run: Dict[str, Any],
    machine_eval: Dict[str, Any],
    generated_at: str,
    model: str,
) -> str:
    metrics = compact_run.get("metrics") or {}
    counts = metrics.get("counts") or {}
    inventory = metrics.get("file_inventory") or {}
    quotes = metrics.get("quote_traceability") or {}
    boundaries = metrics.get("integrated_boundary_metrics") or {}
    criteria = machine_eval.get("criteria") or []

    criterion_prompts = {item["name"]: item.get("human_prompt", "") for item in CRITERIA}
    result_rows = []
    criteria_sorted = sorted(criteria, key=lambda item: _criterion_group_and_number(str(item.get("name") or ""))[1] or 99)
    for criterion in criteria_sorted:
        name = str(criterion.get("name") or "")
        group, number = _criterion_group_and_number(name)
        supporting_metrics = _list_html(_supporting_metrics_for_domain(name, counts, quotes, boundaries)[:3])
        score = criterion.get("score", "")
        result_rows.append(
            "<tr>"
            f"<td>{_esc(group)}</td>"
            f"<td>{_esc(number)}</td>"
            f"<td>{_esc(name)}</td>"
            f"<td><span class=\"score {_score_class(score)}\">{_esc(score)}</span></td>"
            f"<td>{supporting_metrics}</td>"
            f"<td>{_esc(criterion_prompts.get(name, ''))}</td>"
            f"<td>{_esc(criterion.get('rationale', ''))}</td>"
            "</tr>"
        )

    missing_files = inventory.get("missing_required_files") or []
    scale_rows = "".join(
        f"<tr><td><span class=\"score {_score_class(score)}\">{score}</span></td><td>{_esc(text)}</td></tr>"
        for score, text in SCORING_SCALE.items()
    )

    return f"""<!doctype html>
<html lang="en">
<head>
  <meta charset="utf-8" />
  <meta name="viewport" content="width=device-width, initial-scale=1" />
  <title>CCA Evaluation Report</title>
  <style>
    :root {{ --bg:#f6f7fb; --panel:#fff; --border:#d8e0ea; --text:#0f172a; --muted:#64748b; --accent:#2563eb; --green:#166534; --amber:#92400e; --red:#991b1b; }}
    * {{ box-sizing: border-box; }}
    body {{ margin:0; background:var(--bg); color:var(--text); font:14px/1.55 system-ui,-apple-system,Segoe UI,Roboto,Helvetica,Arial,sans-serif; }}
    main {{ max-width:1120px; margin:0 auto; padding:28px 18px 60px; }}
    h1 {{ margin:0 0 4px; font-size:28px; }}
    h2 {{ margin:0 0 12px; font-size:18px; }}
    h3 {{ margin:18px 0 8px; font-size:15px; }}
    p {{ margin:6px 0; }}
    .muted {{ color:var(--muted); }}
    .small {{ font-size:12px; }}
    .panel {{ background:var(--panel); border:1px solid var(--border); border-radius:8px; padding:16px; margin-top:16px; box-shadow:0 2px 10px rgba(15,23,42,.05); }}
    .summary {{ display:grid; grid-template-columns:repeat(auto-fit,minmax(190px,1fr)); gap:10px; margin-top:12px; }}
    .metric-card {{ border:1px solid var(--border); border-radius:8px; padding:12px; background:#fbfdff; min-height:92px; }}
    .metric-label {{ color:var(--muted); font-size:12px; }}
    .metric-value {{ font-size:26px; font-weight:700; margin-top:2px; }}
    .metric-note {{ color:var(--muted); font-size:12px; margin-top:2px; }}
    table {{ width:100%; border-collapse:collapse; margin-top:8px; }}
    th, td {{ border:1px solid var(--border); padding:9px; vertical-align:top; text-align:left; }}
    th {{ background:#eef3f8; font-size:12px; text-transform:uppercase; letter-spacing:.04em; }}
    ul {{ margin:0; padding-left:18px; }}
    .score {{ display:inline-block; min-width:34px; text-align:center; border-radius:999px; padding:2px 8px; font-weight:700; }}
    .score-high {{ background:#dcfce7; color:var(--green); }}
    .score-mid {{ background:#fef3c7; color:var(--amber); }}
    .score-low {{ background:#fee2e2; color:var(--red); }}
    .blank-cell {{ height:68px; background:repeating-linear-gradient(0deg,#fff,#fff 25px,#eef2f7 26px); }}
    .blank-cell.wide {{ min-width:220px; }}
	    @media print {{ body {{ background:#fff; }} .panel {{ box-shadow:none; page-break-inside:avoid; }} }}
	  </style>
</head>
<body>
  <main>
    <header>
      <h1>CCA Evaluation Report</h1>
      <p class="muted">Generated at {_esc(generated_at)} using {_esc(model)}. This paper-ready report evaluates the CCA output itself; it does not use a ground-truth article or human answer.</p>
    </header>

    <section class="panel">
      <h2>Paper-Ready Evaluation Results</h2>
      <p class="muted">Each row is one mutually exclusive evaluation domain. Each domain has one integer 1-5 Likert score and up to three non-overlapping supporting audit metrics.</p>
      <div class="summary">
        {_metric_card("Overall score", machine_eval.get("overall_score", ""), "Mean of six domain scores")}
        {_metric_card("Evaluation domains", len(criteria), "Six mutually exclusive CCA domains")}
        {_metric_card("Missing required files", len(missing_files), ", ".join(missing_files[:3]))}
      </div>
      <table>
        <thead><tr><th>Group</th><th>#</th><th>Evaluation domain</th><th>Main score</th><th>Supporting audit metrics</th><th>Operational definition</th><th>Score rationale</th></tr></thead>
        <tbody>{''.join(result_rows)}</tbody>
      </table>
    </section>

    <section class="panel">
      <h2>1-5 Likert Scoring Scale</h2>
      <p class="muted">Scores are integer-only. Higher scores indicate stronger CCA methodological quality for that domain.</p>
      <table>
        <thead><tr><th>Score</th><th>Anchor</th></tr></thead>
        <tbody>{scale_rows}</tbody>
      </table>
    </section>

	  </main>
</body>
</html>
"""


def generate_evaluation(
    *,
    input_dir: Path,
    output_dir: Path,
    api_key: Optional[str],
    model: Optional[str] = None,
    use_llm: bool = True,
) -> Dict[str, Path]:
    output_dir.mkdir(parents=True, exist_ok=True)
    model = model or getattr(config, "DEFAULT_MODEL", "gpt-5-nano")
    generated_at = datetime.utcnow().isoformat(timespec="seconds") + "Z"
    compact_run = load_cca_run(input_dir)
    machine_eval = run_machine_evaluation(
        compact_run=compact_run,
        api_key=api_key,
        model=model,
        use_llm=use_llm,
    )
    payload = {
        "generated_at": generated_at,
        "model": model,
        "input_dir": str(input_dir),
        "compact_run": compact_run,
        "machine_evaluation": machine_eval,
    }
    machine_json = output_dir / "machine_evaluation.json"
    machine_json.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")
    report_html = output_dir / "evaluation_report.html"
    report_html.write_text(
        render_html_report(
            compact_run=compact_run,
            machine_eval=machine_eval,
            generated_at=generated_at,
            model=model,
        ),
        encoding="utf-8",
    )
    return {"machine_json": machine_json, "report_html": report_html}


def main() -> None:
    parser = argparse.ArgumentParser(description="Generate a standalone HTML evaluation report for completed CCA outputs.")
    parser.add_argument("--input", required=True, help="Folder containing completed CCA output .txt/.json files.")
    parser.add_argument("--output", required=True, help="Folder where evaluation_report.html should be written.")
    parser.add_argument("--model", default=getattr(config, "DEFAULT_MODEL", "gpt-5-nano"))
    parser.add_argument("--api-key", default=os.getenv("OPENAI_API_KEY", ""))
    parser.add_argument("--no-llm", action="store_true", help="Generate deterministic structural checks only.")
    args = parser.parse_args()

    paths = generate_evaluation(
        input_dir=Path(args.input),
        output_dir=Path(args.output),
        api_key=args.api_key.strip() or None,
        model=args.model,
        use_llm=not args.no_llm,
    )
    print(paths["report_html"])


if __name__ == "__main__":
    main()
